"""
Resume loader - extracts text from PDF, DOC, and DOCX files.
"""
from pathlib import Path
from pypdf import PdfReader
from typing import Union
import io

# For DOCX support
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


def load_resume_from_pdf(file_path: Union[str, Path]) -> str:
    """
    Load resume text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content
    """
    reader = PdfReader(file_path)
    text_parts = []
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    
    return "\n".join(text_parts)


def load_resume_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Load resume text from PDF bytes."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    text_parts = []
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    
    return "\n".join(text_parts)


def load_resume_from_docx_bytes(docx_bytes: bytes) -> str:
    """Load resume text from DOCX bytes."""
    if not DOCX_SUPPORT:
        raise ImportError("python-docx is required for DOCX support. Install with: pip install python-docx")
    
    doc = Document(io.BytesIO(docx_bytes))
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    
    return "\n".join(text_parts)


def load_resume_from_bytes(file_bytes: bytes, filename: str = "") -> str:
    """
    Load resume text from file bytes, detecting format from filename or content.
    
    Args:
        file_bytes: Raw file bytes
        filename: Original filename (used for format detection)
        
    Returns:
        Extracted text content
    """
    filename_lower = filename.lower()
    
    # Detect by extension
    if filename_lower.endswith('.docx'):
        return load_resume_from_docx_bytes(file_bytes)
    elif filename_lower.endswith('.doc'):
        # .doc files (old Word format) need different handling
        # python-docx doesn't support .doc, only .docx
        raise ValueError("Old .doc format is not supported. Please convert to .docx or PDF.")
    elif filename_lower.endswith('.pdf'):
        return load_resume_from_pdf_bytes(file_bytes)
    
    # Try to detect by content (magic bytes)
    if file_bytes[:4] == b'%PDF':
        return load_resume_from_pdf_bytes(file_bytes)
    elif file_bytes[:4] == b'PK\x03\x04':  # ZIP signature (DOCX is a ZIP file)
        return load_resume_from_docx_bytes(file_bytes)
    
    # Default to PDF if we can't detect
    return load_resume_from_pdf_bytes(file_bytes)

